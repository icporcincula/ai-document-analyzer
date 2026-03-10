"""
Tests for DOCX Service
"""
import pytest
from pathlib import Path
import tempfile
import docx
from unittest.mock import Mock, patch

from app.services.docx_service import DOCXService
from app.exceptions import DocumentProcessingError


class TestDOCXService:
    """Test cases for DOCX service functionality"""
    
    def setup_method(self):
        """Set up test fixtures"""
        self.service = DOCXService()
        
    def test_validate_docx_file_valid(self):
        """Test validation of valid DOCX file"""
        # Create a temporary DOCX file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            doc = docx.Document()
            doc.add_paragraph("Test paragraph")
            doc.save(temp_file.name)
            
            try:
                # Should not raise an exception
                result = self.service.validate_docx_file(Path(temp_file.name))
                assert result is True
            finally:
                Path(temp_file.name).unlink()
    
    def test_validate_docx_file_invalid_extension(self):
        """Test validation of file with invalid extension"""
        with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as temp_file:
            temp_file.write(b"test content")
            temp_file.flush()
            
            try:
                with pytest.raises(DocumentProcessingError, match="Unsupported DOCX extension"):
                    self.service.validate_docx_file(Path(temp_file.name))
            finally:
                Path(temp_file.name).unlink()
    
    def test_validate_docx_file_too_large(self):
        """Test validation of file that's too large"""
        # Mock the config to have a small max file size
        with patch.object(self.service, 'max_file_size', 100):
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                # Create a file larger than the limit
                temp_file.write(b"x" * 200)
                temp_file.flush()
                
                try:
                    with pytest.raises(DocumentProcessingError, match="DOCX file too large"):
                        self.service.validate_docx_file(Path(temp_file.name))
                finally:
                    Path(temp_file.name).unlink()
    
    def test_extract_text_from_paragraphs(self):
        """Test text extraction from paragraphs"""
        doc = docx.Document()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph")
        doc.add_paragraph("")  # Empty paragraph should be skipped
        
        paragraphs = self.service.extract_text_from_paragraphs(doc)
        
        assert len(paragraphs) == 2
        assert "First paragraph" in paragraphs
        assert "Second paragraph" in paragraphs
    
    def test_extract_text_from_tables(self):
        """Test text extraction from tables"""
        doc = docx.Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Cell 1"
        table.cell(0, 1).text = "Cell 2"
        table.cell(1, 0).text = "Cell 3"
        table.cell(1, 1).text = ""
        
        tables = self.service.extract_text_from_tables(doc)
        
        assert len(tables) == 3
        assert "Cell 1" in tables
        assert "Cell 2" in tables
        assert "Cell 3" in tables
    
    def test_extract_text_from_headers_footers(self):
        """Test text extraction from headers and footers"""
        doc = docx.Document()
        # Note: This is a simplified test as header/footer creation is complex
        headers_footers = self.service.extract_text_from_headers_footers(doc)
        
        # Should return empty list for document without headers/footers
        assert isinstance(headers_footers, list)
    
    def test_process_docx_file(self):
        """Test complete DOCX file processing"""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            doc = docx.Document()
            doc.add_paragraph("Test content")
            doc.add_table(rows=1, cols=1).cell(0, 0).text = "Table content"
            doc.save(temp_file.name)
            
            try:
                result = self.service.process_docx_file(Path(temp_file.name))
                
                assert 'full_text' in result
                assert 'paragraphs' in result
                assert 'tables' in result
                assert 'format' in result
                assert result['format'] == 'docx'
                assert len(result['full_text']) > 0
                
            finally:
                Path(temp_file.name).unlink()
    
    def test_get_docx_metadata(self):
        """Test DOCX metadata extraction"""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            doc = docx.Document()
            doc.core_properties.title = "Test Document"
            doc.core_properties.author = "Test Author"
            doc.save(temp_file.name)
            
            try:
                metadata = self.service.get_docx_metadata(Path(temp_file.name))
                
                assert 'title' in metadata
                assert 'author' in metadata
                assert metadata['title'] == "Test Document"
                assert metadata['author'] == "Test Author"
                
            finally:
                Path(temp_file.name).unlink()
    
    def test_process_empty_docx_file(self):
        """Test processing of empty DOCX file"""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            doc = docx.Document()
            doc.save(temp_file.name)
            
            try:
                result = self.service.process_docx_file(Path(temp_file.name))
                
                # Should still return a result even for empty files
                assert 'full_text' in result
                assert result['full_text'] == ""
                
            finally:
                Path(temp_file.name).unlink()