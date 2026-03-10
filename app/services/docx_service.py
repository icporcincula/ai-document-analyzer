"""
DOCX Service for text extraction from DOCX files.

This service handles DOCX file processing including text extraction from
paragraphs, tables, headers/footers, and formatting preservation.
"""

import logging
from typing import List, Optional, Dict, Any
from pathlib import Path
import docx
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.shared import qn

from app.utils.config import get_config
from app.exceptions import DocumentProcessingError

logger = logging.getLogger(__name__)


class DOCXService:
    """Service for extracting text from DOCX files."""
    
    def __init__(self):
        """Initialize DOCX service with configuration."""
        self.config = get_config()
        self.max_file_size = self.config.docx.max_file_size
        self.supported_extensions = {'.docx', '.doc'}
        
    def validate_docx_file(self, file_path: Path) -> bool:
        """
        Validate DOCX file for processing.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            True if file is valid, False otherwise
            
        Raises:
            DocumentProcessingError: If file validation fails
        """
        try:
            # Check file size
            file_size = file_path.stat().st_size
            if file_size > self.max_file_size:
                raise DocumentProcessingError(
                    f"DOCX file too large: {file_size} bytes (max: {self.max_file_size})"
                )
            
            # Check file extension
            if file_path.suffix.lower() not in self.supported_extensions:
                raise DocumentProcessingError(
                    f"Unsupported DOCX extension: {file_path.suffix}"
                )
            
            # Try to open the document
            try:
                doc = docx.Document(str(file_path))
                # Basic structure validation
                if not doc.paragraphs and not doc.tables:
                    logger.warning(f"DOCX file appears to be empty: {file_path}")
                    return True  # Empty files are technically valid
                    
            except Exception as e:
                raise DocumentProcessingError(f"Failed to open DOCX file: {str(e)}")
                
            return True
            
        except OSError as e:
            raise DocumentProcessingError(f"File system error: {str(e)}")
        except Exception as e:
            raise DocumentProcessingError(f"Unexpected error during DOCX validation: {str(e)}")
    
    def extract_text_from_paragraphs(self, doc: docx.Document) -> List[str]:
        """
        Extract text from all paragraphs in the document.
        
        Args:
            doc: Opened DOCX document
            
        Returns:
            List of paragraph texts
        """
        paragraphs_text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs_text.append(paragraph.text)
        
        return paragraphs_text
    
    def extract_text_from_tables(self, doc: docx.Document) -> List[str]:
        """
        Extract text from all tables in the document.
        
        Args:
            doc: Opened DOCX document
            
        Returns:
            List of table cell texts
        """
        tables_text = []
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        tables_text.append(cell.text)
        
        return tables_text
    
    def extract_text_from_headers_footers(self, doc: docx.Document) -> List[str]:
        """
        Extract text from headers and footers.
        
        Args:
            doc: Opened DOCX document
            
        Returns:
            List of header/footer texts
        """
        headers_footers_text = []
        
        try:
            # Extract from sections
            for section in doc.sections:
                # Header
                if section.header and section.header.is_linked_to_previous is False:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            headers_footers_text.append(f"HEADER: {paragraph.text}")
                
                # Footer  
                if section.footer and section.footer.is_linked_to_previous is False:
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            headers_footers_text.append(f"FOOTER: {paragraph.text}")
                            
        except Exception as e:
            logger.warning(f"Failed to extract headers/footers: {str(e)}")
        
        return headers_footers_text
    
    def extract_text_from_document(self, doc: docx.Document) -> Dict[str, Any]:
        """
        Extract all text from a DOCX document.
        
        Args:
            doc: Opened DOCX document
            
        Returns:
            Dictionary containing extracted text organized by source
        """
        extracted_data = {
            'paragraphs': [],
            'tables': [],
            'headers_footers': [],
            'full_text': ''
        }
        
        # Extract from paragraphs
        extracted_data['paragraphs'] = self.extract_text_from_paragraphs(doc)
        
        # Extract from tables
        extracted_data['tables'] = self.extract_text_from_tables(doc)
        
        # Extract from headers/footers
        extracted_data['headers_footers'] = self.extract_text_from_headers_footers(doc)
        
        # Combine all text
        all_text_parts = (
            extracted_data['paragraphs'] + 
            extracted_data['tables'] + 
            extracted_data['headers_footers']
        )
        extracted_data['full_text'] = '\n'.join(all_text_parts)
        
        return extracted_data
    
    def process_docx_file(self, file_path: Path) -> Dict[str, Any]:
        """
        Process a DOCX file and extract all text content.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dictionary containing extracted text and metadata
            
        Raises:
            DocumentProcessingError: If processing fails
        """
        try:
            logger.info(f"Processing DOCX file: {file_path}")
            
            # Validate file
            self.validate_docx_file(file_path)
            
            # Open document
            doc = docx.Document(str(file_path))
            
            # Extract text
            extracted_data = self.extract_text_from_document(doc)
            
            # Add metadata
            extracted_data.update({
                'file_path': str(file_path),
                'file_size': file_path.stat().st_size,
                'format': 'docx',
                'paragraph_count': len(extracted_data['paragraphs']),
                'table_count': len(doc.tables),
                'total_text_length': len(extracted_data['full_text'])
            })
            
            logger.info(f"Successfully processed DOCX file: {file_path}")
            return extracted_data
            
        except DocumentProcessingError:
            raise
        except Exception as e:
            logger.error(f"Failed to process DOCX file {file_path}: {str(e)}")
            raise DocumentProcessingError(f"DOCX processing failed: {str(e)}")
    
    def get_docx_metadata(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract metadata from DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dictionary containing document metadata
        """
        try:
            doc = docx.Document(str(file_path))
            core_props = doc.core_properties
            
            metadata = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'created': core_props.created.isoformat() if core_props.created else None,
                'modified': core_props.modified.isoformat() if core_props.modified else None,
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
                'section_count': len(doc.sections)
            }
            
            return metadata
            
        except Exception as e:
            logger.warning(f"Failed to extract DOCX metadata: {str(e)}")
            return {}